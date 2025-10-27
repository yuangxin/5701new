"""
CSV和JSON报告输出
"""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Iterable, List


SUMMARY_HEADER = [
    "pair",
    "count",
    "mean_sim",
    "max_sim",
    "coverage_min",
    "coverage_a",
    "coverage_b",
    "student_a_sent_total",
    "student_b_sent_total",
    "score",
]

PARA_SUMMARY_HEADER = [
    "pair",
    "count",
    "mean_sim",
    "max_sim",
    "coverage_min",
    "coverage_a",
    "coverage_b",
    "student_a_para_total",
    "student_b_para_total",
    "score",
]


def write_summary_csv(path: Path, stats: Iterable[dict]) -> None:
    """写句子级汇总CSV"""
    rows_for_csv = []
    for item in stats:
        a, b = item["pair"]
        row = dict(item)
        row["pair"] = f"({a}, {b})"
        rows_for_csv.append(row)

    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=SUMMARY_HEADER)
        writer.writeheader()
        writer.writerows(rows_for_csv)


def write_paragraph_summary(path: Path, stats: Iterable[dict]) -> None:
    """写段落级汇总CSV"""
    rows_for_csv = []
    for item in stats:
        a, b = item["pair"]
        row = dict(item)
        row["pair"] = f"({a}, {b})"
        rows_for_csv.append(row)

    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=PARA_SUMMARY_HEADER)
        writer.writeheader()
        writer.writerows(rows_for_csv)


def write_pair_results(path: Path, details: List[dict]) -> None:
    """写详细结果JSON"""
    payload = {"pairs": details}
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def write_evidence_top(path: Path, details: List[dict]) -> None:
    """写证据映射JSON"""
    evidence_map = {
        str(tuple(detail["pair"])): detail["hits"]
        for detail in details
    }
    path.write_text(
        json.dumps(evidence_map, ensure_ascii=False, indent=2), 
        encoding="utf-8"
    )













from docx import Document
from docx.shared import Pt



def write_word_report(path: Path, stats: List[dict], details: List[dict], top_n_per_pair: int = 5) -> None:

    # ensure parent exists
    path.parent.mkdir(parents=True, exist_ok=True)

    # create document and set default font
    doc = Document()
    style = doc.styles['Normal']
    try:
        style.font.name = 'Arial'
        style.font.size = Pt(11)
    except Exception:
        pass

    doc.add_heading("Plagiarism Detection Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.datetime.now().isoformat()}")
    doc.add_paragraph("")

    # --- Overall summary ---
    total_pairs = len(stats)
    total_matches = sum(item.get("count", 0) for item in stats)
    avg_pair_sim = (sum(item.get("mean_sim", 0.0) for item in stats) / total_pairs) if total_pairs else 0.0
    high_risk = sum(1 for item in stats if item.get("score", 0.0) >= 0.7)
    medium_risk = sum(1 for item in stats if 0.5 <= item.get("score", 0.0) < 0.7)
    low_risk = total_pairs - high_risk - medium_risk

    doc.add_heading("1. Overall Summary", level=2)
    doc.add_paragraph(f"Pairs analyzed: {total_pairs}")
    doc.add_paragraph(f"Total matching sentences found: {total_matches}")
    doc.add_paragraph(f"Average pair mean similarity: {avg_pair_sim:.3f}")
    doc.add_paragraph(f"Risk breakdown: High={high_risk}, Medium={medium_risk}, Low={low_risk}")
    doc.add_paragraph("")

    # optional small summary table
    headers = ["pair", "matches", "mean_sim", "score"]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for item in stats:
        row = table.add_row().cells
        pair = item.get("pair", ("",""))
        row[0].text = f"{pair[0]} ⟷ {pair[1]}"
        row[1].text = str(item.get("count", 0))
        row[2].text = f"{item.get('mean_sim', 0.0):.3f}"
        row[3].text = f"{item.get('score', 0.0):.3f}"
    doc.add_paragraph("")

    # --- Sentence comparisons ---
    doc.add_heading("2. Sentence-level Comparisons", level=2)
    if not details:
        doc.add_paragraph("No detailed pair results available.")
    else:
        for idx, detail in enumerate(details, start=1):
            pair = detail.get("pair", ("?", "?"))
            doc.add_heading(f"Pair {idx}: {pair[0]} ⟷ {pair[1]}", level=3)
            doc.add_paragraph(f"Matches: {detail.get('count', 0)}    Avg sim: {detail.get('mean_sim', 0.0):.3f}    Score: {detail.get('score', 0.0):.3f}")
            hits = detail.get("hits", [])[:top_n_per_pair]
            if not hits:
                doc.add_paragraph("  No sentence-level hits.")
            else:
                for i, hit in enumerate(hits, start=1):
                    sim = hit.get('adjusted_sim', hit.get('sim', 0.0))
                    penalty = hit.get('citation_penalty', 1.0)
                    sid_i = hit.get('sid_i', '')
                    sid_j = hit.get('sid_j', '')
                    sent_i = hit.get('sent_id_i', '')
                    sent_j = hit.get('sent_id_j', '')
                    text_i = hit.get('text_i', '').replace('\n', ' ').strip()
                    text_j = hit.get('text_j', '').replace('\n', ' ').strip()

                    p = doc.add_paragraph()
                    p.add_run(f"{i}. Similarity: {sim:.1%}").bold = True
                    if penalty < 1.0:
                        p.add_run(" (possible citation)").italic = True
                    doc.add_paragraph(f"    {sid_i} (sent {sent_i}): {text_i}")
                    doc.add_paragraph(f"    {sid_j} (sent {sent_j}): {text_j}")
            doc.add_paragraph("")
            # page break between pairs for readability
            doc.add_page_break()

    # save
    doc.save(str(path))